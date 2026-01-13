"""DOCX document generation for reports.

Uses python-docx to create Word documents with:
- Styled headings and paragraphs
- Tables for evidence inventory and witnesses
- Formatted timeline
- Cover page with case metadata
"""

from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..models import Case, Witness, WitnessRole


class DocxGenerator:
    """Generator for DOCX reports."""

    def __init__(self):
        """Initialize with a new document."""
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self) -> None:
        """Configure document styles."""
        # Set default font
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # Configure heading styles
        for i in range(1, 4):
            heading_style = self.doc.styles[f"Heading {i}"]
            heading_style.font.name = "Times New Roman"
            heading_style.font.bold = True

    def add_cover_page(
        self,
        case: Case,
        title: str = "CASE ANALYSIS REPORT",
        prepared_by: Optional[str] = None,
    ) -> None:
        """
        Add a cover page to the document.

        Args:
            case: Case for metadata
            title: Report title
            prepared_by: Preparer name
        """
        # Add some spacing
        for _ in range(3):
            self.doc.add_paragraph()

        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(24)

        self.doc.add_paragraph()

        # Case info
        info_lines = [
            f"Case ID: {case.case_id}",
        ]
        if case.court_case:
            info_lines.append(f"Court Case: {case.court_case}")
        if case.agency:
            info_lines.append(f"Agency: {case.agency}")
        if case.incident_date:
            info_lines.append(f"Incident Date: {case.incident_date.strftime('%B %d, %Y')}")

        for line in info_lines:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            run.font.size = Pt(14)

        # Add spacing
        for _ in range(5):
            self.doc.add_paragraph()

        # Prepared info
        if prepared_by:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(f"Prepared by: {prepared_by}")

        # Date
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(f"Report Date: {datetime.now().strftime('%B %d, %Y')}")

        # Page break
        self.doc.add_page_break()

    def add_heading(self, text: str, level: int = 1) -> None:
        """Add a heading."""
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False) -> None:
        """Add a paragraph."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic

    def add_bullet_list(self, items: list[str]) -> None:
        """Add a bulleted list."""
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")

    def add_numbered_list(self, items: list[str]) -> None:
        """Add a numbered list."""
        for item in items:
            self.doc.add_paragraph(item, style="List Number")

    def add_table(
        self,
        headers: list[str],
        rows: list[list[str]],
        col_widths: Optional[list[float]] = None,
    ) -> None:
        """
        Add a table to the document.

        Args:
            headers: Column headers
            rows: Table data rows
            col_widths: Optional column widths in inches
        """
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Bold header text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                row.cells[i].text = str(cell_text)

        # Set column widths if provided
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)

        self.doc.add_paragraph()  # Space after table

    def add_evidence_inventory_table(self, case: Case) -> None:
        """Add evidence inventory as a table."""
        headers = ["#", "Filename", "Type", "Size", "Duration", "Status"]
        rows = []

        for i, evidence in enumerate(case.evidence_items, 1):
            # Format size
            size = evidence.size_bytes
            if size < 1024 * 1024:
                size_str = f"{size / 1024:.1f} KB"
            else:
                size_str = f"{size / (1024 * 1024):.1f} MB"

            # Format duration
            if evidence.duration_seconds:
                mins = int(evidence.duration_seconds // 60)
                secs = int(evidence.duration_seconds % 60)
                duration_str = f"{mins}:{secs:02d}"
            else:
                duration_str = "-"

            # Status
            status_parts = []
            if evidence.transcript:
                status_parts.append("T")
            if evidence.summary:
                status_parts.append("S")
            status_str = ", ".join(status_parts) if status_parts else "-"

            rows.append([
                str(i),
                evidence.filename[:40] + ("..." if len(evidence.filename) > 40 else ""),
                evidence.file_type.value,
                size_str,
                duration_str,
                status_str,
            ])

        self.add_table(headers, rows, col_widths=[0.4, 2.5, 0.8, 0.8, 0.7, 0.6])

    def add_witness_table(self, case: Case) -> None:
        """Add witness list as a table."""
        if not case.witnesses.witnesses:
            self.add_paragraph("No witnesses identified.")
            return

        headers = ["Name", "Role", "Speaker IDs", "Evidence Count", "Verified"]
        rows = []

        for witness in case.witnesses.witnesses:
            rows.append([
                witness.display_name,
                witness.role.value,
                ", ".join(witness.speaker_ids[:2]) + ("..." if len(witness.speaker_ids) > 2 else ""),
                str(len(witness.evidence_appearances)),
                "Yes" if witness.verified else "No",
            ])

        self.add_table(headers, rows, col_widths=[1.8, 1.0, 1.2, 1.0, 0.8])

    def add_timeline(self, case: Case) -> None:
        """Add timeline section."""
        if not case.timeline:
            self.add_paragraph("No timeline events.")
            return

        current_date = None

        for event in case.timeline:
            # Add date header when date changes
            event_date = event.timestamp.date()
            if event_date != current_date:
                current_date = event_date
                self.add_paragraph(
                    event_date.strftime("%B %d, %Y"),
                    bold=True,
                )

            # Format event
            time_str = event.timestamp.strftime("%H:%M:%S")
            importance = event.metadata.get("importance", "normal") if event.metadata else "normal"

            prefix = ""
            if importance == "critical":
                prefix = "[!!!] "
            elif importance == "high":
                prefix = "[!] "

            self.add_paragraph(f"  {time_str}  {prefix}{event.description}")

    def add_key_statements(self, case: Case) -> None:
        """Add key statements section."""
        key_statements = []

        for evidence in case.evidence_items:
            if not evidence.transcript:
                continue

            for segment in evidence.transcript.key_statements:
                key_statements.append({
                    "evidence": evidence.filename,
                    "speaker": segment.speaker,
                    "text": segment.text,
                    "translation": segment.translation,
                    "flags": segment.review_flags,
                })

        if not key_statements:
            self.add_paragraph("No key statements flagged.")
            return

        for stmt in key_statements:
            # Statement header
            para = self.doc.add_paragraph()
            run = para.add_run(f"[{stmt['evidence']}] ")
            run.bold = True
            run.font.size = Pt(10)

            if stmt["speaker"]:
                run = para.add_run(f"{stmt['speaker']}: ")
                run.italic = True

            # Statement text
            para.add_run(stmt["text"])

            # Translation if available
            if stmt["translation"]:
                trans_para = self.doc.add_paragraph()
                trans_para.paragraph_format.left_indent = Inches(0.5)
                run = trans_para.add_run(f"[Translation: {stmt['translation']}]")
                run.italic = True
                run.font.size = Pt(10)

    def add_summary_section(self, case: Case) -> None:
        """Add executive summary section."""
        if case.summary:
            self.add_paragraph(case.summary)
        else:
            self.add_paragraph("No summary available.")

        # Add key findings if available
        if case.metadata and case.metadata.get("key_findings"):
            self.add_heading("Key Findings", level=2)
            self.add_numbered_list(case.metadata["key_findings"])

    def add_review_items(self, case: Case) -> None:
        """Add items needing review section."""
        if not case.metadata or not case.metadata.get("items_needing_review"):
            self.add_paragraph("No items flagged for review.")
            return

        self.add_bullet_list(case.metadata["items_needing_review"])

    def save(self, path: Path) -> None:
        """
        Save the document to a file.

        Args:
            path: Output file path
        """
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def generate_case_docx(
    case: Case,
    output_path: Path,
    prepared_by: Optional[str] = None,
    include_transcripts: bool = False,
) -> Path:
    """
    Generate a complete case report in DOCX format.

    Args:
        case: Case to report on
        output_path: Where to save the document
        prepared_by: Preparer name for cover page
        include_transcripts: Whether to include full transcripts

    Returns:
        Path to generated document
    """
    gen = DocxGenerator()

    # Cover page
    gen.add_cover_page(case, prepared_by=prepared_by)

    # Table of contents placeholder
    gen.add_heading("Table of Contents", level=1)
    gen.add_paragraph("[Table of Contents - update field after opening in Word]")
    gen.doc.add_page_break()

    # Executive summary
    gen.add_heading("Executive Summary", level=1)
    gen.add_summary_section(case)
    gen.doc.add_page_break()

    # Case overview
    gen.add_heading("Case Overview", level=1)
    gen.add_paragraph(f"Case ID: {case.case_id}", bold=True)
    if case.court_case:
        gen.add_paragraph(f"Court Case: {case.court_case}")
    if case.agency:
        gen.add_paragraph(f"Agency: {case.agency}")
    if case.incident_date:
        gen.add_paragraph(f"Incident Date: {case.incident_date.strftime('%B %d, %Y')}")
    gen.add_paragraph(f"Total Evidence Items: {len(case.evidence_items)}")
    gen.add_paragraph(f"Witnesses Identified: {len(case.witnesses.witnesses)}")

    # Statistics
    if case.metadata:
        gen.add_heading("Processing Statistics", level=2)
        if case.metadata.get("threats_identified"):
            gen.add_paragraph(f"Threat Keywords Detected: {case.metadata['threats_identified']}")
        if case.metadata.get("violence_indicators"):
            gen.add_paragraph(f"Violence Indicators: {case.metadata['violence_indicators']}")
        if case.metadata.get("spanish_percentage"):
            gen.add_paragraph(f"Spanish Content: {case.metadata['spanish_percentage']:.1f}%")

    gen.doc.add_page_break()

    # Evidence inventory
    gen.add_heading("Evidence Inventory", level=1)
    gen.add_evidence_inventory_table(case)
    gen.doc.add_page_break()

    # Witnesses
    gen.add_heading("Witnesses", level=1)
    gen.add_witness_table(case)
    gen.doc.add_page_break()

    # Timeline
    gen.add_heading("Timeline of Events", level=1)
    gen.add_timeline(case)
    gen.doc.add_page_break()

    # Key statements
    gen.add_heading("Key Statements", level=1)
    gen.add_key_statements(case)
    gen.doc.add_page_break()

    # Items needing review
    gen.add_heading("Items Requiring Review", level=1)
    gen.add_review_items(case)

    # Full transcripts if requested
    if include_transcripts:
        gen.doc.add_page_break()
        gen.add_heading("Full Transcripts", level=1)

        for evidence in case.evidence_items:
            if not evidence.transcript:
                continue

            gen.add_heading(evidence.filename, level=2)

            for segment in evidence.transcript.segments:
                speaker = evidence.transcript.get_speaker_name(segment.speaker)
                time_str = f"[{segment.start_time:.1f}s]"

                para = gen.doc.add_paragraph()
                run = para.add_run(f"{time_str} {speaker}: ")
                run.bold = True
                para.add_run(segment.text)

                if segment.translation:
                    trans_para = gen.doc.add_paragraph()
                    trans_para.paragraph_format.left_indent = Inches(0.5)
                    run = trans_para.add_run(f"[{segment.translation}]")
                    run.italic = True

    # Save
    gen.save(output_path)

    return output_path
