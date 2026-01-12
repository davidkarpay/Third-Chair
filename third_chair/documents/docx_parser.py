"""Word document (DOCX) parsing.

Extracts text, tables, and metadata from Word documents.
Includes special handling for Axon transcript format.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional


@dataclass
class DocxParagraph:
    """Represents a paragraph from a Word document."""

    text: str
    style: Optional[str] = None
    is_heading: bool = False
    heading_level: int = 0
    is_list_item: bool = False


@dataclass
class DocxTable:
    """Represents a table from a Word document."""

    rows: list[list[str]]
    header_row: Optional[list[str]] = None

    @property
    def row_count(self) -> int:
        return len(self.rows)

    @property
    def col_count(self) -> int:
        return len(self.rows[0]) if self.rows else 0

    def to_dicts(self) -> list[dict]:
        """Convert table to list of dicts using header row as keys."""
        if not self.header_row or not self.rows:
            return []

        result = []
        for row in self.rows:
            row_dict = {}
            for i, cell in enumerate(row):
                if i < len(self.header_row):
                    key = self.header_row[i].strip().lower().replace(" ", "_")
                    row_dict[key] = cell
            result.append(row_dict)

        return result


@dataclass
class DocxDocument:
    """Represents a parsed Word document."""

    file_path: Path
    paragraphs: list[DocxParagraph] = field(default_factory=list)
    tables: list[DocxTable] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        """Get all paragraph text."""
        return "\n".join(p.text for p in self.paragraphs if p.text)

    @property
    def headings(self) -> list[DocxParagraph]:
        """Get all headings."""
        return [p for p in self.paragraphs if p.is_heading]

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())


def parse_docx(file_path: Path) -> DocxDocument:
    """
    Parse a Word document.

    Args:
        file_path: Path to the DOCX file

    Returns:
        DocxDocument with extracted content
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    doc = Document(file_path)
    result = DocxDocument(file_path=file_path)

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else None
        is_heading = style_name and "Heading" in style_name
        heading_level = 0

        if is_heading:
            # Extract heading level
            level_match = re.search(r"Heading\s*(\d+)", style_name)
            if level_match:
                heading_level = int(level_match.group(1))

        is_list = style_name and "List" in style_name

        result.paragraphs.append(DocxParagraph(
            text=text,
            style=style_name,
            is_heading=is_heading,
            heading_level=heading_level,
            is_list_item=is_list,
        ))

    # Extract tables
    for table in doc.tables:
        rows = []
        header_row = None

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]

            if i == 0:
                # Check if first row looks like a header
                if _looks_like_header(cells):
                    header_row = cells
                    continue

            rows.append(cells)

        if rows:
            result.tables.append(DocxTable(
                rows=rows,
                header_row=header_row,
            ))

    # Extract metadata
    result.metadata = _extract_docx_metadata(doc)

    return result


def _looks_like_header(cells: list[str]) -> bool:
    """Check if a row looks like a table header."""
    if not cells:
        return False

    # Headers typically have short, label-like text
    header_patterns = [
        "name", "date", "time", "speaker", "text", "timestamp",
        "id", "type", "description", "notes", "role",
    ]

    matches = 0
    for cell in cells:
        cell_lower = cell.lower()
        if any(pattern in cell_lower for pattern in header_patterns):
            matches += 1

    # If more than half match header patterns, it's likely a header
    return matches >= len(cells) / 2


def _extract_docx_metadata(doc) -> dict:
    """Extract metadata from a Word document."""
    metadata = {}

    try:
        core_props = doc.core_properties
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.created:
            metadata["created"] = core_props.created.isoformat()
        if core_props.modified:
            metadata["modified"] = core_props.modified.isoformat()
        if core_props.subject:
            metadata["subject"] = core_props.subject
    except Exception:
        pass

    return metadata


def extract_text_from_docx(file_path: Path) -> str:
    """
    Simple function to extract all text from a Word document.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text as a single string
    """
    doc = parse_docx(file_path)
    return doc.full_text


# ============================================================================
# Axon Transcript Parsing
# ============================================================================

@dataclass
class AxonTranscriptSegment:
    """A segment from an Axon transcript."""

    timestamp: str
    speaker: str
    text: str
    raw_time_seconds: Optional[float] = None


@dataclass
class AxonTranscript:
    """Parsed Axon transcript document."""

    file_path: Path
    segments: list[AxonTranscriptSegment] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def segment_count(self) -> int:
        return len(self.segments)

    @property
    def speakers(self) -> list[str]:
        return list(set(s.speaker for s in self.segments))

    @property
    def duration_seconds(self) -> Optional[float]:
        if not self.segments:
            return None
        last_time = self.segments[-1].raw_time_seconds
        return last_time


def parse_axon_transcript(file_path: Path) -> AxonTranscript:
    """
    Parse an Axon transcript Word document.

    Axon transcripts have a specific format:
    - Timestamps like [10:15 AM / 00:15]
    - Speaker labels like **Speaker 1:**
    - Text following each speaker

    Args:
        file_path: Path to Axon transcript DOCX

    Returns:
        AxonTranscript with parsed segments
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    file_path = Path(file_path)
    doc = Document(file_path)

    result = AxonTranscript(file_path=file_path)

    # Regex patterns
    timestamp_pattern = re.compile(r"\[(\d{1,2}:\d{2}\s*[AP]M\s*/\s*[\d:]+)\]")
    speaker_pattern = re.compile(r"\*?\*?Speaker\s*(\d+)\*?\*?:", re.IGNORECASE)

    current_timestamp = None
    current_speaker = None
    current_text_parts = []
    in_metadata = True

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check for timestamp
        ts_match = timestamp_pattern.search(text)
        if ts_match:
            # Save previous segment
            if current_timestamp and current_text_parts:
                segment_text = " ".join(current_text_parts).strip()
                raw_seconds = _parse_axon_timestamp(current_timestamp)

                result.segments.append(AxonTranscriptSegment(
                    timestamp=current_timestamp,
                    speaker=current_speaker or "Unknown",
                    text=segment_text,
                    raw_time_seconds=raw_seconds,
                ))

            in_metadata = False
            current_timestamp = ts_match.group(1)
            current_text_parts = []

            # Check for speaker in same line
            remaining = text[ts_match.end():].strip()
            sp_match = speaker_pattern.search(remaining)
            if sp_match:
                current_speaker = f"SPEAKER_{sp_match.group(1)}"
                after_speaker = remaining[sp_match.end():].strip().strip("*").strip()
                if after_speaker:
                    current_text_parts.append(after_speaker)
            continue

        # Check for speaker change
        sp_match = speaker_pattern.search(text)
        if sp_match and not in_metadata:
            current_speaker = f"SPEAKER_{sp_match.group(1)}"
            after_speaker = text[sp_match.end():].strip().strip("*").strip()
            if after_speaker:
                current_text_parts.append(after_speaker)
            continue

        # Extract metadata from header
        if in_metadata:
            if text.startswith("**Evidence Title:"):
                result.metadata["title"] = text.replace("**Evidence Title:", "").strip().strip("*")
            elif "transcript starts at" in text.lower():
                result.metadata["start_info"] = text
            continue

        # Regular text line
        if current_timestamp:
            clean_text = text.strip("*").strip()
            if clean_text:
                current_text_parts.append(clean_text)

    # Don't forget the last segment
    if current_timestamp and current_text_parts:
        segment_text = " ".join(current_text_parts).strip()
        raw_seconds = _parse_axon_timestamp(current_timestamp)

        result.segments.append(AxonTranscriptSegment(
            timestamp=current_timestamp,
            speaker=current_speaker or "Unknown",
            text=segment_text,
            raw_time_seconds=raw_seconds,
        ))

    return result


def _parse_axon_timestamp(timestamp: str) -> Optional[float]:
    """
    Parse Axon timestamp to seconds.

    Format: "10:15 AM / 00:15" or "10:15:30 AM / 00:15:30"
    We use the second part (relative time).
    """
    # Extract the relative time part (after /)
    if "/" in timestamp:
        rel_part = timestamp.split("/")[1].strip()
    else:
        rel_part = timestamp.strip()

    # Parse MM:SS or HH:MM:SS
    parts = rel_part.split(":")
    try:
        if len(parts) == 2:
            minutes, seconds = map(int, parts)
            return minutes * 60 + seconds
        elif len(parts) == 3:
            hours, minutes, seconds = map(int, parts)
            return hours * 3600 + minutes * 60 + seconds
    except ValueError:
        pass

    return None


def is_axon_transcript(file_path: Path) -> bool:
    """
    Check if a Word document is an Axon transcript.

    Args:
        file_path: Path to check

    Returns:
        True if file appears to be an Axon transcript
    """
    try:
        from docx import Document

        doc = Document(file_path)

        # Check first few paragraphs for Axon patterns
        text = " ".join(p.text for p in doc.paragraphs[:10])

        patterns = [
            r"\[\d{1,2}:\d{2}\s*[AP]M\s*/",  # Axon timestamp
            r"Speaker\s*\d+:",  # Speaker label
            r"Evidence Title:",  # Axon metadata
        ]

        matches = sum(1 for p in patterns if re.search(p, text, re.IGNORECASE))
        return matches >= 2

    except Exception:
        return False
