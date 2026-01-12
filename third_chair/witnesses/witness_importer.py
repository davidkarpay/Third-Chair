"""Import witness lists from various file formats.

Supports:
- PDF (text-based via pdfplumber)
- PDF (scanned via pytesseract OCR)
- Word documents (via python-docx)
- Excel spreadsheets (via openpyxl)
- Plain text
"""

import re
from pathlib import Path
from typing import Optional

from ..models import Witness, WitnessList, WitnessRole, WitnessSource


def import_witness_list(
    source: str | Path,
    source_type: Optional[str] = None,
) -> WitnessList:
    """
    Import a witness list from various sources.

    Args:
        source: File path or text content
        source_type: Type of source ("pdf", "docx", "xlsx", "text")
                    If None, auto-detected from file extension

    Returns:
        WitnessList with extracted witnesses
    """
    # Handle direct text input
    if isinstance(source, str) and not Path(source).exists():
        return _parse_witness_text(source)

    path = Path(source)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    # Auto-detect type from extension
    if source_type is None:
        ext = path.suffix.lower()
        source_type = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "docx",
            ".xlsx": "xlsx",
            ".xls": "xlsx",
            ".txt": "text",
            ".csv": "csv",
        }.get(ext, "text")

    # Import based on type
    if source_type == "pdf":
        return _import_from_pdf(path)
    elif source_type == "docx":
        return _import_from_docx(path)
    elif source_type == "xlsx":
        return _import_from_excel(path)
    elif source_type == "csv":
        return _import_from_csv(path)
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return _parse_witness_text(text)


def _import_from_pdf(path: Path) -> WitnessList:
    """Import witnesses from a PDF file."""
    text = ""

    # Try text extraction first
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except ImportError:
        pass
    except Exception as e:
        print(f"Warning: pdfplumber failed: {e}")

    # If no text extracted, try OCR
    if not text.strip():
        text = _ocr_pdf(path)

    return _parse_witness_text(text)


def _ocr_pdf(path: Path) -> str:
    """OCR a scanned PDF."""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        # Convert PDF pages to images
        images = convert_from_path(path)

        text_parts = []
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image)
            text_parts.append(page_text)

        return "\n".join(text_parts)

    except ImportError as e:
        print(f"Warning: OCR dependencies not available: {e}")
        return ""
    except Exception as e:
        print(f"Warning: OCR failed: {e}")
        return ""


def _import_from_docx(path: Path) -> WitnessList:
    """Import witnesses from a Word document."""
    try:
        from docx import Document

        doc = Document(path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        text = "\n".join(text_parts)
        return _parse_witness_text(text)

    except ImportError:
        raise ImportError("python-docx is required for Word documents. Install with: pip install python-docx")


def _import_from_excel(path: Path) -> WitnessList:
    """Import witnesses from an Excel spreadsheet."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb.active

        if ws is None:
            return WitnessList()

        witnesses = WitnessList()
        headers = []

        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            # First row is headers
            if row_idx == 0:
                headers = [str(c).lower() if c else f"col_{i}" for i, c in enumerate(row)]
                continue

            # Create witness from row
            row_dict = {
                headers[i]: str(val) if val else ""
                for i, val in enumerate(row)
                if i < len(headers)
            }

            witness = _row_to_witness(row_dict)
            if witness:
                witnesses.add(witness)

        wb.close()
        return witnesses

    except ImportError:
        raise ImportError("openpyxl is required for Excel files. Install with: pip install openpyxl")


def _import_from_csv(path: Path) -> WitnessList:
    """Import witnesses from a CSV file."""
    import csv

    witnesses = WitnessList()

    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)

        for row in reader:
            # Normalize keys
            row_dict = {k.lower().strip(): v for k, v in row.items()}
            witness = _row_to_witness(row_dict)
            if witness:
                witnesses.add(witness)

    return witnesses


def _row_to_witness(row: dict) -> Optional[Witness]:
    """Convert a row dict to a Witness object."""
    # Find name field
    name = None
    name_fields = ["name", "witness name", "full name", "witness", "nombre"]
    for field in name_fields:
        if field in row and row[field].strip():
            name = row[field].strip()
            break

    # Skip empty rows
    if not name:
        return None

    # Detect role
    role = WitnessRole.OTHER
    role_fields = ["role", "type", "witness type", "category"]
    for field in role_fields:
        if field in row and row[field].strip():
            role_str = row[field].lower().strip()
            if "victim" in role_str:
                role = WitnessRole.VICTIM
            elif "witness" in role_str:
                role = WitnessRole.WITNESS
            elif "suspect" in role_str or "defendant" in role_str:
                role = WitnessRole.SUSPECT
            elif "officer" in role_str or "deputy" in role_str:
                role = WitnessRole.OFFICER
            break

    # Get contact info
    contact = None
    contact_fields = ["phone", "contact", "address", "email"]
    contact_parts = []
    for field in contact_fields:
        if field in row and row[field].strip():
            contact_parts.append(row[field].strip())
    if contact_parts:
        contact = "; ".join(contact_parts)

    # Get notes
    notes = None
    notes_fields = ["notes", "comments", "description"]
    for field in notes_fields:
        if field in row and row[field].strip():
            notes = row[field].strip()
            break

    return Witness(
        name=name,
        role=role,
        source=WitnessSource.STATE_ATTORNEY_LIST,
        contact_info=contact,
        notes=notes,
    )


def _parse_witness_text(text: str) -> WitnessList:
    """
    Parse witness information from free-form text.

    Handles various formats:
    - Numbered lists: "1. John Smith - Victim"
    - Bulleted lists: "• Jane Doe (Witness)"
    - Name/role pairs: "John Smith, Victim"
    - State Attorney witness list formats
    """
    witnesses = WitnessList()

    if not text.strip():
        return witnesses

    lines = text.strip().split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Skip headers and section titles
        if _is_header_line(line):
            continue

        # Try to parse as witness entry
        witness = _parse_witness_line(line)
        if witness:
            witnesses.add(witness)

    return witnesses


def _is_header_line(line: str) -> bool:
    """Check if a line is a header/title rather than a witness entry."""
    headers = [
        "witness list", "witnesses", "state witness", "defense witness",
        "victim", "victims", "name", "role", "contact",
        "state of florida", "case number", "defendant",
    ]
    line_lower = line.lower().strip()

    # Short lines that are likely headers
    if len(line_lower) < 4:
        return True

    # Lines that match header patterns
    for header in headers:
        if line_lower == header or line_lower.startswith(header + ":"):
            return True

    # All caps lines are often headers
    if line.isupper() and len(line) > 5:
        return True

    return False


def _parse_witness_line(line: str) -> Optional[Witness]:
    """Parse a single line as a witness entry."""
    # Remove list markers
    line = re.sub(r"^[\d]+[.)\]]\s*", "", line)  # Numbered: 1. 2) 3]
    line = re.sub(r"^[-•*]\s*", "", line)  # Bulleted
    line = line.strip()

    if not line or len(line) < 3:
        return None

    # Pattern: "Name - Role" or "Name (Role)"
    patterns = [
        r"^([^-–—]+)\s*[-–—]\s*(victim|witness|suspect|officer|defendant)",
        r"^([^(]+)\s*\((victim|witness|suspect|officer|defendant)\)",
        r"^(victim|witness):\s*(.+)",
    ]

    for pattern in patterns:
        match = re.match(pattern, line, re.IGNORECASE)
        if match:
            groups = match.groups()
            if "victim" in pattern or "witness" in pattern:
                if groups[0].lower() in ("victim", "witness"):
                    # Pattern 3: role first
                    role_str, name = groups
                else:
                    name, role_str = groups
            else:
                name, role_str = groups

            role = _parse_role(role_str)
            return Witness(
                name=name.strip(),
                role=role,
                source=WitnessSource.STATE_ATTORNEY_LIST,
            )

    # Pattern: Just a name (assume witness)
    # Must look like a person's name (2+ words, starts with capital)
    if re.match(r"^[A-Z][a-z]+\s+[A-Z]", line):
        # Extract just the name part (before any punctuation)
        name_match = re.match(r"^([A-Za-z\s\-']+)", line)
        if name_match:
            name = name_match.group(1).strip()
            # Must have at least 2 name parts
            if len(name.split()) >= 2:
                return Witness(
                    name=name,
                    role=WitnessRole.WITNESS,
                    source=WitnessSource.STATE_ATTORNEY_LIST,
                )

    return None


def _parse_role(role_str: str) -> WitnessRole:
    """Parse a role string to WitnessRole enum."""
    role_lower = role_str.lower().strip()

    if "victim" in role_lower:
        return WitnessRole.VICTIM
    elif "witness" in role_lower:
        return WitnessRole.WITNESS
    elif "suspect" in role_lower or "defendant" in role_lower:
        return WitnessRole.SUSPECT
    elif "officer" in role_lower or "deputy" in role_lower:
        return WitnessRole.OFFICER
    elif "interpreter" in role_lower:
        return WitnessRole.INTERPRETER

    return WitnessRole.OTHER
